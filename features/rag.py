"""
features/rag.py
===============

ParselyFi "Document Brain" tab — Phase-3 grounded chat + table Q&A over the
user's own uploaded files.

Public contract (the ONLY name the main app imports):

    render_rag_tab() -> None

Design rules baked in (see project CLAUDE.md / the spec):

- IMPORT-SAFE: every heavy/optional dependency (PyMuPDF ``fitz``,
  ``python-docx``, ``fastembed``, ``qdrant_client``, ``scikit-learn``) is
  imported inside try/except and reflected through ``*_AVAILABLE`` flags, so
  ``import features.rag`` ALWAYS succeeds even on a bare install. We report
  which resolved in the UI.
- HONEST STATUS: missing GEMINI key -> inject_css + hero + health banner +
  return. Missing retrieval libs -> degrade to lexical-only (sklearn TF-IDF)
  and SAY SO; if no retrieval libs at all, still offer table Q&A and a clear
  notice. Never a success-shaped failure.
- HONEST DATA: chat answers ONLY from retrieved chunks and cites
  ``filename#chunk``; if nothing relevant is retrieved we say
  "I could not find this in your documents" instead of hallucinating. Table
  Q&A passes a BOUNDED, honest slice of the dataframe and never eval()s
  model-produced code.
- BOUNDED MEMORY: MAX_FILES, MAX_FILE_BYTES, TOTAL_CHUNKS, MAX_TURNS,
  TABLE_ROW_CAP — every accumulating collection is capped with an honest
  truncation notice. File reads are size-capped.
- ERROR_BOUNDARY: one bad file (or one bad chunk) never kills the batch; the
  failing file surfaces as an error row.
- Reuses ``features.common`` for ALL LLM/token logic and ``features.ui`` for
  the shared look. Every session key is uniquely prefixed ``rag_``.

Retrieval modes (honest):
- "semantic" (HYBRID): fastembed BGE-small dense vectors in an in-memory
  Qdrant collection, fused via Reciprocal Rank Fusion (RRF) with a sklearn
  TF-IDF cosine lexical score over the same chunks.
- "lexical": sklearn TF-IDF cosine only (when fastembed/qdrant are absent).
- "none": no retrieval libs at all -> semantic search disabled, table Q&A
  still works.
"""

from __future__ import annotations

import hashlib
import io
import logging
import re
from typing import Any, Dict, List, Optional, Tuple

import streamlit as st

from . import common
from . import ui
from .common import (
    GEMINI_MODEL,
    feature_available,
    gemini_generate_text,
    push_history,
    render_gemini_health,
    render_token_usage,
    run_async,
)

logger = logging.getLogger("parselyfi.features.rag")

# ===========================================================================
# Optional dependencies — guarded so `import features.rag` NEVER crashes.
# ===========================================================================
try:
    import pandas as pd  # type: ignore
    PANDAS_AVAILABLE = True
except Exception:  # pragma: no cover - pandas underpins the table-Q&A UI
    pd = None  # type: ignore
    PANDAS_AVAILABLE = False

try:
    import numpy as np  # type: ignore
    NUMPY_AVAILABLE = True
except Exception:  # pragma: no cover
    np = None  # type: ignore
    NUMPY_AVAILABLE = False

try:
    import fitz  # PyMuPDF  # type: ignore
    FITZ_AVAILABLE = True
except Exception:
    fitz = None  # type: ignore
    FITZ_AVAILABLE = False

try:
    import docx  # python-docx  # type: ignore
    DOCX_AVAILABLE = True
except Exception:
    docx = None  # type: ignore
    DOCX_AVAILABLE = False

try:
    from sklearn.feature_extraction.text import TfidfVectorizer  # type: ignore
    from sklearn.metrics.pairwise import cosine_similarity  # type: ignore
    SKLEARN_AVAILABLE = True
except Exception:
    TfidfVectorizer = None  # type: ignore
    cosine_similarity = None  # type: ignore
    SKLEARN_AVAILABLE = False

try:
    from fastembed import TextEmbedding  # type: ignore
    FASTEMBED_AVAILABLE = True
except Exception:
    TextEmbedding = None  # type: ignore
    FASTEMBED_AVAILABLE = False

try:
    from qdrant_client import QdrantClient  # type: ignore
    from qdrant_client.models import (  # type: ignore
        Distance,
        PointStruct,
        VectorParams,
    )
    QDRANT_AVAILABLE = True
except Exception:
    QdrantClient = None  # type: ignore
    Distance = None  # type: ignore
    PointStruct = None  # type: ignore
    VectorParams = None  # type: ignore
    QDRANT_AVAILABLE = False

# Semantic (hybrid) retrieval needs ALL THREE of fastembed + qdrant + sklearn.
SEMANTIC_AVAILABLE = FASTEMBED_AVAILABLE and QDRANT_AVAILABLE and SKLEARN_AVAILABLE
# Lexical fallback needs sklearn only.
LEXICAL_AVAILABLE = SKLEARN_AVAILABLE


# ===========================================================================
# Constants / bounds (every accumulating collection is capped).
# ===========================================================================
SS_PREFIX = "rag_"
SS_DOCS = SS_PREFIX + "docs"               # list[ {file_id, name, kind, n_chunks, error} ]
SS_CHUNKS = SS_PREFIX + "chunks"           # list[ {text, source, chunk_idx, global_idx} ]
SS_HASHES = SS_PREFIX + "file_hashes"      # set of sha256 already ingested
SS_TABLES = SS_PREFIX + "tables"           # {name: {"df": DataFrame}}
SS_INDEX = SS_PREFIX + "index"             # cached retrieval index handle
SS_MODE = SS_PREFIX + "mode"               # "semantic" | "lexical" | "none"
SS_CHAT = SS_PREFIX + "chat_history"       # bounded via common.push_history
SS_TRUNCATED = SS_PREFIX + "chunks_truncated"

MAX_FILES = 12
MAX_FILE_BYTES = 15 * 1024 * 1024          # ~15MB per file
CHUNK_SIZE = 900                           # ~900 chars per chunk
CHUNK_OVERLAP = 150                        # ~150 char overlap
TOTAL_CHUNKS = 3000                        # hard cap across all docs
MAX_TURNS = 20                             # bounded chat history (user+assistant pairs)
TOP_K = 6                                  # chunks retrieved per question
RRF_K = 60                                 # Reciprocal Rank Fusion constant
TABLE_ROW_CAP = 50                         # rows of a table sent to Gemini as context
MAX_TABLE_PROMPT_CHARS = 40_000            # bounded read on the CSV slice we send
PER_CHUNK_PREVIEW_CHARS = 600              # bounded snippet length shown in the UI
MAX_TEXT_CHARS_PER_FILE = 4_000_000        # bounded read on extracted text per file

_EMBED_MODEL_NAME = "BAAI/bge-small-en-v1.5"
_EMBED_DIM = 384                            # bge-small-en-v1.5 dimensionality
_QDRANT_COLLECTION = "rag_chunks"

# Process-global fastembed model cache (loading the model is expensive; the
# model object itself holds no per-user state so it is safe to share).
_EMBEDDER_SINGLETON: Any = None


# ===========================================================================
# Text extraction (one bad file never kills the batch).
# ===========================================================================

def _extract_pdf(data: bytes) -> str:
    """Extract text from a PDF via PyMuPDF. Bounded read."""
    if not FITZ_AVAILABLE or fitz is None:
        raise RuntimeError("PyMuPDF (fitz) not installed")
    parts: List[str] = []
    total = 0
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            txt = page.get_text("text") or ""
            parts.append(txt)
            total += len(txt)
            if total > MAX_TEXT_CHARS_PER_FILE:
                parts.append("\n[truncated: document exceeded read cap]")
                break
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    """Extract text from a .docx via python-docx (paragraphs + table cells)."""
    if not DOCX_AVAILABLE or docx is None:
        raise RuntimeError("python-docx not installed")
    document = docx.Document(io.BytesIO(data))
    parts: List[str] = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells if c.text]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    return text[:MAX_TEXT_CHARS_PER_FILE]


def _extract_plain(data: bytes) -> str:
    """Decode a txt/md file defensively, bounded."""
    text = data.decode("utf-8", errors="replace")
    return text[:MAX_TEXT_CHARS_PER_FILE]


def _read_dataframe(data: bytes, suffix: str):
    """Read a csv/xlsx/xls into a pandas DataFrame (bounded later by the caller)."""
    if not PANDAS_AVAILABLE or pd is None:
        raise RuntimeError("pandas not installed")
    bio = io.BytesIO(data)
    if suffix == "csv":
        return pd.read_csv(bio)
    # xlsx / xls -> openpyxl / xlrd engine chosen by pandas.
    return pd.read_excel(bio)


def _df_to_text(df) -> str:
    """Flatten a DataFrame to searchable text (header + bounded rows) for the
    semantic/lexical index. Honest: only the first TABLE_ROW_CAP*4 rows are
    folded into the chunk index so a giant sheet cannot blow the chunk budget;
    the FULL frame remains queryable via the table-Q&A feature."""
    if not PANDAS_AVAILABLE or pd is None:
        return ""
    head = df.head(TABLE_ROW_CAP * 4)
    try:
        return head.to_csv(index=False)
    except Exception:
        return head.to_string(index=False)


def _suffix_of(name: str) -> str:
    return name.rsplit(".", 1)[-1].lower() if "." in name else ""


# ===========================================================================
# Chunking (bounded, with overlap).
# ===========================================================================

def _chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Split text into bounded ~``size``-char chunks with ``overlap`` overlap.

    Splits on a sliding window over normalized whitespace. Never produces an
    empty chunk; ``overlap`` is clamped below ``size`` to guarantee progress.
    """
    text = re.sub(r"[ \t]+", " ", (text or "")).strip()
    if not text:
        return []
    size = max(1, int(size))
    overlap = max(0, min(int(overlap), size - 1))
    step = size - overlap
    chunks: List[str] = []
    i = 0
    n = len(text)
    while i < n:
        piece = text[i:i + size].strip()
        if piece:
            chunks.append(piece)
        i += step
    return chunks


# ===========================================================================
# Retrieval index (semantic hybrid via RRF, or lexical fallback).
# ===========================================================================

def _get_embedder():
    """Return a process-cached fastembed TextEmbedding (BGE-small) or None."""
    global _EMBEDDER_SINGLETON
    if not FASTEMBED_AVAILABLE or TextEmbedding is None:
        return None
    if _EMBEDDER_SINGLETON is None:
        try:
            _EMBEDDER_SINGLETON = TextEmbedding(model_name=_EMBED_MODEL_NAME)
            logger.info("fastembed model loaded: %s", _EMBED_MODEL_NAME)
        except Exception as e:
            logger.error("fastembed model load failed: %s", e)
            return None
    return _EMBEDDER_SINGLETON


class _RagIndex:
    """In-memory hybrid (or lexical-only) retrieval index over the chunks.

    Built once per ingest. Holds the chunk texts plus, when available, a
    sklearn TF-IDF matrix and an in-memory Qdrant collection of dense vectors.
    All resources live for the lifetime of this object (an in-memory Qdrant
    client), and the index is rebuilt from scratch on each ingest so nothing
    accumulates across reruns.
    """

    def __init__(self, chunks: List[Dict[str, Any]]):
        self.chunks = chunks
        self.mode = "none"
        self._tfidf = None
        self._tfidf_matrix = None
        self._qdrant = None
        self._embedder = None
        self._build()

    # -- build --------------------------------------------------------------
    def _build(self) -> None:
        if not self.chunks:
            return
        texts = [c["text"] for c in self.chunks]

        # Lexical (TF-IDF) — needed for both modes; lexical-only fallback too.
        if LEXICAL_AVAILABLE:
            try:
                self._tfidf = TfidfVectorizer(
                    stop_words="english", max_features=50_000, ngram_range=(1, 2)
                )
                self._tfidf_matrix = self._tfidf.fit_transform(texts)
                self.mode = "lexical"
            except Exception as e:
                logger.error("TF-IDF build failed: %s", e)
                self._tfidf = None
                self._tfidf_matrix = None

        # Dense (fastembed + qdrant) — upgrades mode to semantic/hybrid.
        if SEMANTIC_AVAILABLE and self._tfidf_matrix is not None:
            try:
                embedder = _get_embedder()
                if embedder is not None:
                    vectors = [list(map(float, v)) for v in embedder.embed(texts)]
                    client = QdrantClient(location=":memory:")
                    # create_collection is the current API (recreate_collection is
                    # deprecated); a fresh in-memory client always starts empty.
                    client.create_collection(
                        collection_name=_QDRANT_COLLECTION,
                        vectors_config=VectorParams(
                            size=len(vectors[0]) if vectors else _EMBED_DIM,
                            distance=Distance.COSINE,
                        ),
                    )
                    points = [
                        PointStruct(id=idx, vector=vec, payload={"chunk_idx": idx})
                        for idx, vec in enumerate(vectors)
                    ]
                    client.upsert(collection_name=_QDRANT_COLLECTION, points=points)
                    self._qdrant = client
                    self._embedder = embedder
                    self.mode = "semantic"
            except Exception as e:
                logger.error("Dense index build failed; staying lexical: %s", e)
                self._qdrant = None
                self._embedder = None

    # -- retrieval ----------------------------------------------------------
    def _lexical_ranking(self, query: str, top_n: int) -> List[Tuple[int, float]]:
        """Return [(chunk_idx, score)] by TF-IDF cosine, descending."""
        if self._tfidf is None or self._tfidf_matrix is None:
            return []
        try:
            qv = self._tfidf.transform([query])
            sims = cosine_similarity(qv, self._tfidf_matrix)[0]
            order = sims.argsort()[::-1][:top_n]
            return [(int(i), float(sims[i])) for i in order if sims[i] > 0.0]
        except Exception as e:
            logger.error("Lexical ranking failed: %s", e)
            return []

    def _dense_ranking(self, query: str, top_n: int) -> List[Tuple[int, float]]:
        """Return [(chunk_idx, score)] by Qdrant dense cosine, descending."""
        if self._qdrant is None or self._embedder is None:
            return []
        try:
            qvec = list(map(float, next(iter(self._embedder.embed([query])))))
            hits = self._qdrant.search(
                collection_name=_QDRANT_COLLECTION,
                query_vector=qvec,
                limit=top_n,
            )
            return [(int(h.payload.get("chunk_idx", h.id)), float(h.score)) for h in hits]
        except Exception as e:
            logger.error("Dense ranking failed: %s", e)
            return []

    @staticmethod
    def _rrf_fuse(
        rankings: List[List[Tuple[int, float]]], k: int = RRF_K
    ) -> List[int]:
        """Reciprocal Rank Fusion across multiple ranked lists.

        Score(d) = sum over lists of 1 / (k + rank_in_list(d)). Returns chunk
        ids sorted by fused score descending.
        """
        fused: Dict[int, float] = {}
        for ranked in rankings:
            for rank, (idx, _score) in enumerate(ranked):
                fused[idx] = fused.get(idx, 0.0) + 1.0 / (k + rank + 1)
        return [idx for idx, _ in sorted(fused.items(), key=lambda kv: kv[1], reverse=True)]

    def retrieve(self, query: str, top_k: int = TOP_K) -> List[Dict[str, Any]]:
        """Retrieve the top-``top_k`` chunks for ``query``.

        Semantic mode fuses dense + lexical rankings via RRF; lexical mode uses
        TF-IDF cosine only. Returns [] when nothing scores above zero so the
        caller can honestly say "not found".
        """
        if not self.chunks or not (query or "").strip():
            return []
        pool = max(top_k * 4, top_k)
        if self.mode == "semantic":
            dense = self._dense_ranking(query, pool)
            lexical = self._lexical_ranking(query, pool)
            if not dense and not lexical:
                return []
            fused_ids = self._rrf_fuse([dense, lexical])
        else:  # lexical (or degraded)
            lexical = self._lexical_ranking(query, pool)
            if not lexical:
                return []
            fused_ids = [idx for idx, _ in lexical]

        out: List[Dict[str, Any]] = []
        for cid in fused_ids[:top_k]:
            if 0 <= cid < len(self.chunks):
                out.append(self.chunks[cid])
        return out


def _build_index() -> Optional[_RagIndex]:
    """(Re)build the retrieval index from the session chunks; cache on session."""
    chunks = st.session_state.get(SS_CHUNKS) or []
    if not chunks:
        st.session_state[SS_INDEX] = None
        st.session_state[SS_MODE] = "none"
        return None
    index = _RagIndex(chunks)
    st.session_state[SS_INDEX] = index
    st.session_state[SS_MODE] = index.mode
    return index


# ===========================================================================
# Ingest pipeline (bounded, error-isolated per file, sha256 dedupe).
# ===========================================================================

def _ingest_files(uploaded_files) -> None:
    """Extract, chunk, dedupe, and index uploaded files. Bounded + isolated.

    Never raises: each file is processed in its own try/except and a failing
    file is surfaced as an error row. Honest truncation notices when caps hit.
    """
    docs: List[Dict[str, Any]] = []
    chunks: List[Dict[str, Any]] = []
    tables: Dict[str, Dict[str, Any]] = {}
    seen_hashes: set = set()
    chunks_truncated = False

    # BOUND: never process more than MAX_FILES.
    files = list(uploaded_files or [])
    files_skipped_count = 0
    if len(files) > MAX_FILES:
        files_skipped_count = len(files) - MAX_FILES
        files = files[:MAX_FILES]

    global_idx = 0

    with st.status("Ingesting documents…", expanded=True) as status:
        for uf in files:
            name = getattr(uf, "name", "file")
            suffix = _suffix_of(name)
            row: Dict[str, Any] = {
                "file": name, "kind": suffix, "chunks": 0, "status": "ok", "note": ""
            }
            try:
                data = uf.getvalue()
            except Exception as e:
                row["status"] = "error"
                row["note"] = f"could not read upload: {e}"
                docs.append(row)
                st.write(f"❌ {name}: {row['note']}")
                continue

            # BOUND: per-file size cap (skip + notice).
            if len(data) > MAX_FILE_BYTES:
                row["status"] = "skipped"
                row["note"] = f"larger than {MAX_FILE_BYTES // (1024 * 1024)}MB cap"
                docs.append(row)
                st.write(f"⚠️ {name}: skipped ({row['note']})")
                continue

            # SHA-256 dedupe (skip identical files).
            digest = hashlib.sha256(data).hexdigest()
            if digest in seen_hashes:
                row["status"] = "deduped"
                row["note"] = "identical to an already-ingested file"
                docs.append(row)
                st.write(f"↩️ {name}: deduped")
                continue
            seen_hashes.add(digest)

            # Per-file extraction, fully isolated.
            try:
                if suffix == "pdf":
                    text = _extract_pdf(data)
                elif suffix == "docx":
                    text = _extract_docx(data)
                elif suffix in ("txt", "md"):
                    text = _extract_plain(data)
                elif suffix in ("csv", "xlsx", "xls"):
                    df = _read_dataframe(data, suffix)
                    # BOUND the queryable table itself.
                    if len(df) > 0 and len(tables) < MAX_FILES:
                        tables[name] = {"df": df}
                    text = _df_to_text(df)
                else:
                    row["status"] = "skipped"
                    row["note"] = f"unsupported type .{suffix}"
                    docs.append(row)
                    st.write(f"⚠️ {name}: {row['note']}")
                    continue
            except Exception as e:
                row["status"] = "error"
                row["note"] = f"extract failed: {e}"
                docs.append(row)
                st.write(f"❌ {name}: {row['note']}")
                logger.exception("Extraction failed for %s", name)
                continue

            # Chunk + append (bounded by TOTAL_CHUNKS globally).
            file_chunks = _chunk_text(text)
            added = 0
            for ci, piece in enumerate(file_chunks):
                if global_idx >= TOTAL_CHUNKS:
                    chunks_truncated = True
                    break
                chunks.append(
                    {
                        "text": piece,
                        "source": name,
                        "chunk_idx": ci,
                        "global_idx": global_idx,
                    }
                )
                global_idx += 1
                added += 1
            row["chunks"] = added
            if added == 0 and row["status"] == "ok":
                row["note"] = "no extractable text"
            docs.append(row)
            st.write(f"✅ {name}: {added} chunk(s)")
            if chunks_truncated:
                break

        # Persist results.
        st.session_state[SS_DOCS] = docs
        st.session_state[SS_CHUNKS] = chunks
        st.session_state[SS_HASHES] = seen_hashes
        st.session_state[SS_TABLES] = tables
        st.session_state[SS_TRUNCATED] = chunks_truncated

        status.update(label="Building retrieval index…", state="running")
        index = _build_index()
        mode = index.mode if index else "none"

        status.update(
            label=f"Ingest complete — {len(chunks)} chunk(s), mode: {mode}",
            state="complete",
            expanded=False,
        )

    if files_skipped_count > 0:
        st.warning(
            f"You uploaded more than {MAX_FILES} files; the extra "
            f"{files_skipped_count} were ignored (file cap)."
        )
    if chunks_truncated:
        st.warning(
            f"Chunk budget reached: indexed the first {TOTAL_CHUNKS:,} chunks "
            "only. Later files / pages were not indexed."
        )
    st.toast(f"Ingested {len(chunks)} chunks from {len([d for d in docs if d['status'] == 'ok'])} file(s).")


# ===========================================================================
# Grounded chat (answers ONLY from retrieved chunks, cites filename#chunk).
# ===========================================================================

def _build_grounded_prompt(question: str, retrieved: List[Dict[str, Any]]) -> str:
    """Build a strictly-grounded prompt with numbered, cited context blocks."""
    blocks = []
    for r in retrieved:
        cite = f"{r['source']}#chunk{r['chunk_idx']}"
        blocks.append(f"[Source: {cite}]\n{r['text']}")
    context = "\n\n---\n\n".join(blocks)
    return (
        "You are a careful document-grounded assistant. Answer the user's "
        "question USING ONLY the numbered source excerpts below. Do not use any "
        "outside knowledge.\n"
        "- Cite the source of every claim inline using the exact format "
        "(filename#chunkN) drawn from the [Source: ...] headers.\n"
        "- If the excerpts do not contain the answer, reply EXACTLY: "
        '"I could not find this in your documents." and nothing else.\n'
        "- Do not invent file names, numbers, or citations.\n\n"
        f"=== SOURCE EXCERPTS ===\n{context}\n\n"
        f"=== QUESTION ===\n{question}\n\n"
        "=== ANSWER (grounded, with inline (filename#chunkN) citations) ==="
    )


_NOT_FOUND = "I could not find this in your documents."


def _answer_question(question: str, index: _RagIndex) -> Tuple[str, List[Dict[str, Any]]]:
    """Retrieve + answer. Returns (answer, retrieved_chunks). Honest on miss."""
    retrieved = index.retrieve(question, top_k=TOP_K)
    if not retrieved:
        return _NOT_FOUND, []
    prompt = _build_grounded_prompt(question, retrieved)
    try:
        answer = run_async(
            gemini_generate_text(prompt, agent_name="rag_chat")
        )
    except Exception as e:  # ERROR_BOUNDARY: never let an LLM error crash the tab
        logger.error("rag_chat generation failed: %s", e)
        answer = ""
    if not answer:
        # Honest: the model returned nothing (key/timeout/error). Do NOT fake.
        return (
            "_The assistant returned no answer (model unavailable or timed out). "
            "Check Gemini health above._",
            retrieved,
        )
    return answer, retrieved


def _render_chat(index: Optional[_RagIndex]) -> None:
    """Render the bounded chat history + chat_input + grounded answer."""
    ui.section("Grounded chat", "Answers come ONLY from your documents, with citations.")

    history: List[Dict[str, Any]] = st.session_state.get(SS_CHAT) or []

    # Replay bounded history.
    for turn in history:
        with st.chat_message(turn.get("role", "user")):
            st.markdown(turn.get("content", ""))
            sources = turn.get("sources") or []
            if sources:
                with st.expander(f"Retrieved sources ({len(sources)})"):
                    for s in sources:
                        st.markdown(f"**{s['source']}#chunk{s['chunk_idx']}**")
                        st.caption(s["text"][:PER_CHUNK_PREVIEW_CHARS])

    if index is None:
        st.info(
            "Upload documents above and build the index to enable grounded chat. "
            "(Tables can be queried below without the chunk index.)"
        )
        return

    prompt = st.chat_input("Ask a question about your documents…")
    if not prompt:
        return

    # Show + persist the user turn (bounded).
    with st.chat_message("user"):
        st.markdown(prompt)
    push_history(SS_CHAT, {"role": "user", "content": prompt})

    with st.chat_message("assistant"):
        with st.spinner("Searching your documents…"):
            answer, retrieved = _answer_question(prompt, index)
        st.markdown(answer)
        if retrieved:
            with st.expander(f"Retrieved sources ({len(retrieved)})"):
                for s in retrieved:
                    st.markdown(f"**{s['source']}#chunk{s['chunk_idx']}**")
                    st.caption(s["text"][:PER_CHUNK_PREVIEW_CHARS])

    push_history(
        SS_CHAT,
        {
            "role": "assistant",
            "content": answer,
            "sources": [
                {"source": r["source"], "chunk_idx": r["chunk_idx"], "text": r["text"]}
                for r in retrieved
            ],
        },
    )

    # BOUND: keep at most MAX_TURNS*2 messages (user+assistant per turn).
    hist = st.session_state.get(SS_CHAT) or []
    if len(hist) > MAX_TURNS * 2:
        st.session_state[SS_CHAT] = hist[-(MAX_TURNS * 2):]


# ===========================================================================
# Table Q&A (bounded slice -> Gemini; NEVER eval() model code).
# ===========================================================================

def _build_table_prompt(name: str, df, question: str) -> str:
    """Build a grounded prompt over a BOUNDED slice of a table."""
    shape = f"{df.shape[0]} rows x {df.shape[1]} columns"
    try:
        dtypes = "\n".join(f"- {col}: {dt}" for col, dt in df.dtypes.astype(str).items())
    except Exception:
        dtypes = "(dtypes unavailable)"
    head = df.head(TABLE_ROW_CAP)
    try:
        csv_slice = head.to_csv(index=False)
    except Exception:
        csv_slice = head.to_string(index=False)
    if len(csv_slice) > MAX_TABLE_PROMPT_CHARS:
        csv_slice = csv_slice[:MAX_TABLE_PROMPT_CHARS] + "\n[slice truncated to char cap]"
    return (
        "You are a careful data analyst. Answer the question using ONLY the "
        "table data provided below. Do not invent rows or numbers.\n"
        f"- The full table has {shape}; you are given the FIRST {min(TABLE_ROW_CAP, df.shape[0])} "
        "rows only. If answering requires rows beyond this slice, say so honestly.\n"
        "- Show the computation you did over the visible rows.\n\n"
        f"=== TABLE: {name} ({shape}) ===\n"
        f"Column types:\n{dtypes}\n\n"
        f"First rows (CSV):\n{csv_slice}\n\n"
        f"=== QUESTION ===\n{question}\n\n=== ANSWER ==="
    )


def _answer_table(name: str, df, question: str) -> str:
    """Answer a table question over a bounded slice. Honest on empty/error."""
    prompt = _build_table_prompt(name, df, question)
    try:
        answer = run_async(
            gemini_generate_text(prompt, agent_name=f"rag_table:{name}")
        )
    except Exception as e:  # ERROR_BOUNDARY
        logger.error("rag_table generation failed for %s: %s", name, e)
        answer = ""
    if not answer:
        return (
            "_The assistant returned no answer (model unavailable or timed out)._"
        )
    return answer


def _render_table_qa() -> None:
    """Render each ingested table + a bounded, grounded Q&A box per table."""
    tables: Dict[str, Dict[str, Any]] = st.session_state.get(SS_TABLES) or {}
    if not tables:
        return

    ui.section(
        "Table Q&A",
        "Query each spreadsheet directly. Answers use a bounded row slice — never executed code.",
    )

    for name, payload in tables.items():
        df = payload.get("df")
        if df is None:
            continue
        with st.container(border=True):
            st.markdown(f"**{name}** — {df.shape[0]:,} rows x {df.shape[1]} columns")
            st.dataframe(df, use_container_width=True, height=260)
            if df.shape[0] > TABLE_ROW_CAP:
                st.caption(
                    f"Q&A sees the first {TABLE_ROW_CAP} rows only "
                    "(bounded context sent to the model)."
                )
            q = st.text_input(
                "Ask about this table",
                key=SS_PREFIX + "tq_" + hashlib.sha256(name.encode()).hexdigest()[:10],
                placeholder="e.g. Which row has the highest value in column X?",
            )
            if q:
                with st.spinner("Analyzing the table slice…"):
                    ans = _answer_table(name, df, q)
                st.markdown(ans)


# ===========================================================================
# Cross-tab handoff: pre-fill List Intelligence names text_area.
# ===========================================================================

def _send_to_list_intelligence() -> None:
    """Extract document/source names + sheet column hints and pre-fill the
    List Intelligence names text_area (its exact widget key is
    ``li_raw_text_input``). Honest: we only forward what we actually have."""
    candidates: List[str] = []

    # Use Gemini (if available) to pull company names from the indexed text;
    # fall back to the document file stems so the handoff still does something.
    docs = st.session_state.get(SS_DOCS) or []
    chunks = st.session_state.get(SS_CHUNKS) or []

    names: List[str] = []
    if chunks:
        sample = "\n\n".join(c["text"] for c in chunks[:12])[:12_000]
        prompt = (
            "Extract every distinct COMPANY / ORGANIZATION name mentioned in the "
            "text below. Return ONLY the names, one per line, no numbering, no "
            "commentary. If none, return nothing.\n\n" + sample
        )
        try:
            raw = run_async(gemini_generate_text(prompt, agent_name="rag_handoff_extract"))
        except Exception as e:
            logger.error("handoff extraction failed: %s", e)
            raw = ""
        for line in (raw or "").splitlines():
            nm = line.strip(" \t-•*0123456789.").strip()
            if nm and nm.lower() not in {n.lower() for n in names}:
                names.append(nm)

    candidates = names
    if not candidates:
        # Honest fallback: forward document file stems (deduped).
        seen = set()
        for d in docs:
            stem = d.get("file", "").rsplit(".", 1)[0].strip()
            if stem and stem.lower() not in seen:
                seen.add(stem.lower())
                candidates.append(stem)

    if not candidates:
        st.toast("Nothing to send — no names found in your documents.")
        return

    # Write to the EXACT widget key the List Intelligence text_area reads.
    st.session_state["li_raw_text_input"] = "\n".join(candidates[:200])
    st.toast(f"Sent {len(candidates[:200])} name(s) to List Intelligence. Open that tab to run them.")


# ===========================================================================
# Public entry point.
# ===========================================================================

def render_rag_tab() -> None:
    """Render the "Document Brain" tab: grounded chat + table Q&A over uploads."""
    ui.inject_css()

    mode_label = (
        "Hybrid semantic + lexical (RRF)" if SEMANTIC_AVAILABLE
        else "Lexical only (TF-IDF)" if LEXICAL_AVAILABLE
        else "Table Q&A only"
    )
    ui.hero(
        "Document Brain",
        "Grounded chat and table Q&A over the files you upload — cited, bounded, honest.",
        chips=[
            f"Retrieval: <b>{mode_label}</b>",
            f"Files cap: <b>{MAX_FILES}</b>",
            f"Chunks cap: <b>{TOTAL_CHUNKS:,}</b>",
        ],
    )

    # HONEST STATUS GATE on GEMINI_API_KEY — never a success-shaped failure.
    ok, _missing = feature_available(["GEMINI_API_KEY"])
    if not ok:
        render_gemini_health()
        st.warning(
            "Document Brain needs Gemini to answer questions. Set `GEMINI_API_KEY` "
            "in secrets. (You can still upload files, but chat / table Q&A are disabled.)"
        )
        return
    # Surface a leaked/invalid-key banner if a prior call was rejected.
    render_gemini_health()

    # Honest dependency notice.
    if not SEMANTIC_AVAILABLE:
        missing = []
        if not FASTEMBED_AVAILABLE:
            missing.append("fastembed")
        if not QDRANT_AVAILABLE:
            missing.append("qdrant-client")
        if not SKLEARN_AVAILABLE:
            missing.append("scikit-learn")
        if LEXICAL_AVAILABLE:
            st.info(
                "Semantic search is OFF (missing: " + ", ".join(missing) + "). "
                "Falling back to **lexical TF-IDF retrieval** — keyword matching, "
                "not meaning-based. Install fastembed + qdrant-client for semantic search."
            )
        else:
            st.warning(
                "No retrieval libraries available (need scikit-learn for lexical, "
                "plus fastembed + qdrant-client for semantic). Document chat is "
                "disabled; **table Q&A still works** below."
            )

    # -- Stage 1: ingest ----------------------------------------------------
    ui.section("Upload documents", f"PDF, DOCX, TXT, MD, CSV, XLSX/XLS — up to {MAX_FILES} files, {MAX_FILE_BYTES // (1024*1024)}MB each.")

    # Honest extractor notice.
    extractor_notes = []
    if not FITZ_AVAILABLE:
        extractor_notes.append("PDF extraction unavailable (PyMuPDF missing)")
    if not DOCX_AVAILABLE:
        extractor_notes.append("DOCX extraction unavailable (python-docx missing)")
    if not PANDAS_AVAILABLE:
        extractor_notes.append("CSV/Excel ingest unavailable (pandas missing)")
    if extractor_notes:
        st.caption("⚠️ " + "; ".join(extractor_notes) + ".")

    uploaded = st.file_uploader(
        "Drop files here",
        accept_multiple_files=True,
        type=["pdf", "docx", "txt", "md", "csv", "xlsx", "xls"],
        key=SS_PREFIX + "uploader",
    )

    col_a, col_b = st.columns([1, 1])
    with col_a:
        do_ingest = st.button(
            "Ingest & index",
            type="primary",
            disabled=not uploaded,
            key=SS_PREFIX + "ingest_btn",
            use_container_width=True,
        )
    with col_b:
        if st.button(
            "Clear all",
            key=SS_PREFIX + "clear_btn",
            use_container_width=True,
        ):
            for k in (SS_DOCS, SS_CHUNKS, SS_HASHES, SS_TABLES, SS_INDEX, SS_MODE, SS_CHAT, SS_TRUNCATED):
                st.session_state.pop(k, None)
            st.toast("Cleared documents and chat.")
            st.rerun()

    if do_ingest and uploaded:
        _ingest_files(uploaded)

    # -- KPI row (files, chunks, mode) -------------------------------------
    docs = st.session_state.get(SS_DOCS) or []
    chunks = st.session_state.get(SS_CHUNKS) or []
    tables = st.session_state.get(SS_TABLES) or {}
    mode = st.session_state.get(SS_MODE, "none")
    ok_files = len([d for d in docs if d.get("status") == "ok"])

    if docs:
        ui.kpi_row(
            [
                ("Files indexed", ok_files),
                ("Chunks", f"{len(chunks):,}"),
                ("Retrieval mode", {"semantic": "Semantic", "lexical": "Lexical", "none": "None"}.get(mode, mode)),
            ]
        )

        # Per-file ingest report (honest: errors/skips/dedupes shown).
        if PANDAS_AVAILABLE:
            with st.expander("Ingest report", expanded=any(d.get("status") != "ok" for d in docs)):
                st.dataframe(pd.DataFrame(docs), use_container_width=True, hide_index=True)

        # Handoff to List Intelligence.
        if st.button(
            "Send extracted names to List Intelligence",
            key=SS_PREFIX + "handoff_btn",
        ):
            with st.spinner("Extracting company/organization names…"):
                _send_to_list_intelligence()

    # -- Stage 2: grounded chat --------------------------------------------
    index = st.session_state.get(SS_INDEX)
    # Rebuild index if chunks exist but the index object was lost (e.g. fresh rerun
    # after a cache clear) — index objects don't always survive every code path.
    if index is None and chunks:
        index = _build_index()

    # Chat only makes sense when we have a retrieval index.
    if chunks and (SEMANTIC_AVAILABLE or LEXICAL_AVAILABLE):
        _render_chat(index)
    elif chunks and not LEXICAL_AVAILABLE:
        ui.section("Grounded chat", "")
        st.warning(
            "Document chat needs scikit-learn (lexical) and ideally fastembed + "
            "qdrant-client (semantic). Install them to enable grounded chat. "
            "Table Q&A below still works."
        )

    # -- Stage 3: table Q&A -------------------------------------------------
    if tables:
        _render_table_qa()

    # -- Token usage --------------------------------------------------------
    st.divider()
    render_token_usage()


__all__ = ["render_rag_tab"]
